import os
import re
import json
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml.shared import OxmlElement, qn

ROOT = os.path.dirname(os.path.dirname(__file__))
ASSETS_DIR = os.path.join(ROOT, "assets")
BACKEND_FILE = os.path.join(ROOT, "backend.py")
SIMPLE_BACKEND_FILE = os.path.join(ROOT, "simple_backend.py")
CONFIG_FILES = [
    os.path.join(ROOT, "falconeye-config.yml"),
    os.path.join(ROOT, "falconeye-permanent-config.yml"),
    os.path.join(ROOT, "working-cloudflare-config.yml"),
    os.path.join(ROOT, "cloudflare-config.yml"),
]


def apply_tnr(paragraph, size=12, bold=False, align=None):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rPr.append(rFonts)
        run.font.size = Pt(size)
        run.bold = bold
    if not paragraph.runs:
        run = paragraph.add_run("")
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
    if align is not None:
        paragraph.alignment = align
    # 1.5 line spacing
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), "360")  # 240 = single, 360 = 1.5
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)


def add_heading(doc, text, level=1, center=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    apply_tnr(p, size=14 if level == 1 else 12, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER if center else None)
    return p


def add_page_break(doc):
    doc.add_page_break()


def add_field(doc, field_instr):
    p = doc.add_paragraph()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), field_instr)
    p._p.append(fld)
    apply_tnr(p)
    return p


def add_caption(doc, label, text):
    p = doc.add_paragraph(f"{label}: {text}")
    try:
        p.style = 'Caption'
    except Exception:
        # If 'Caption' style not present, proceed with default
        pass
    apply_tnr(p, size=11)
    return p


def extract_routes(py_path):
    if not os.path.exists(py_path):
        return []
    with open(py_path, 'r', encoding='utf-8', errors='ignore') as f:
        code = f.read()
    routes = re.findall(r"@app\.route\(\"([^\"]+)\"(?:,\s*methods=\[([^\]]+)\])?\)", code)
    out = []
    for path, methods in routes:
        methods_list = []
        if methods:
            methods_list = [m.strip().strip("'\"") for m in methods.split(',')]
        else:
            methods_list = ["GET"]
        out.append({"path": path, "methods": methods_list})
    return out


def read_text_if_exists(path):
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception:
        return None


def add_multiline(doc, text):
    for line in text.splitlines():
        p = doc.add_paragraph(line)
        apply_tnr(p)


def add_code_block(doc, text):
    for line in text.splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        p.paragraph_format.left_indent = Inches(0.25)


def maybe_add_image(doc, filename, caption):
    path = os.path.join(ASSETS_DIR, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(6.0))
        add_caption(doc, "Figure", caption)


def build_document():
    doc = Document()

    # Cover page
    add_heading(doc, "FALCONEYE: SMART SURVEILLANCE WITH AI VISION", level=1, center=True)
    p = doc.add_paragraph("Project Report submitted in partial fulfillment of the requirements")
    apply_tnr(p)
    p = doc.add_paragraph("for the award of Bachelor of Engineering")
    apply_tnr(p)
    p = doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}")
    apply_tnr(p)
    add_page_break(doc)

    # Certificate & Declaration (placeholders)
    add_heading(doc, "Certificate", level=1)
    add_multiline(doc, "This is to certify that the project work titled 'FalconEye' has been carried out under my supervision and guidance.")
    add_page_break(doc)

    add_heading(doc, "Declaration", level=1)
    add_multiline(doc, "We hereby declare that the work presented in this report is our original contribution and has not been submitted elsewhere.")
    add_page_break(doc)

    # Abstract
    add_heading(doc, "Abstract", level=1)
    abstract = (
        "FalconEye is an edge-first smart surveillance system combining dual cameras (ESP32 and Raspberry Pi Zero MJPEG) "
        "with on-device YOLOv8 inference, secure remote access via Cloudflare Tunnel, and optional push notifications. "
        "It delivers low-latency detection of persons, faces, and common objects, includes a PTZ pan controller, and "
        "offers a responsive dashboard optimized for mobile browsers."
    )
    add_multiline(doc, abstract)
    add_page_break(doc)

    # Acknowledgement
    add_heading(doc, "Acknowledgement", level=1)
    add_multiline(doc, "We thank mentors and peers for their valuable guidance throughout the development of FalconEye.")
    add_page_break(doc)

    # Abbreviations
    add_heading(doc, "List of Abbreviations", level=1)
    abbrev = [
        ("YOLO", "You Only Look Once (object detection family)"),
        ("FPS", "Frames Per Second"),
        ("MJPEG", "Motion JPEG video stream"),
        ("PTZ", "Pan-Tilt-Zoom"),
        ("FCM", "Firebase Cloud Messaging"),
        ("ESP32", "Espressif ESP32 microcontroller"),
        ("MPS", "Metal Performance Shaders (Apple GPU acceleration)"),
        ("CORS", "Cross-Origin Resource Sharing"),
    ]
    for k, v in abbrev:
        p = doc.add_paragraph(f"{k} — {v}")
        apply_tnr(p)
    add_page_break(doc)

    # TOC/LoF/LoT placeholders
    add_heading(doc, "Table of Contents", level=1)
    add_field(doc, 'TOC \\o "1-3" \\h \\z \\u')
    add_page_break(doc)

    add_heading(doc, "List of Figures", level=1)
    add_field(doc, 'TOC \\h \\z \\c "Figure"')
    add_page_break(doc)

    add_heading(doc, "List of Tables", level=1)
    add_field(doc, 'TOC \\h \\z \\c "Table"')
    add_page_break(doc)

    # Chapter 1: Introduction
    add_heading(doc, "Chapter 1: Introduction", level=1)
    intro = (
        "This report documents the design and implementation of FalconEye, an AI-powered surveillance system. "
        "It captures frames from an ESP32 snapshot endpoint and a Raspberry Pi Zero MJPEG stream, runs YOLOv8s on Apple Silicon (MPS) where available, and overlays rich annotations and face summaries. "
        "Remote access is implemented using Cloudflare Tunnel, avoiding inbound port exposure."
    )
    add_multiline(doc, intro)

    # Architecture figure
    maybe_add_image(doc, "architecture.png", "System architecture with dual cameras and Cloudflare Tunnel")
    add_page_break(doc)

    # Chapter 2: Literature Survey / Related Work
    add_heading(doc, "Chapter 2: Related Work", level=1)
    add_multiline(doc, (
        "We review classical CCTV systems, edge AI deployments with ESP32-class devices, and modern object detectors (YOLOv5/YOLOv8). "
        "Prior art explores MJPEG streaming on SBCs, serverless secure tunneling (Cloudflare), and notification pipelines via FCM."
    ))
    add_page_break(doc)

    # Chapter 3: System Design
    add_heading(doc, "Chapter 3: System Design", level=1)
    add_multiline(doc, (
        "FalconEye follows an Input–Process–Output paradigm: inputs from ESP32 and Pi Zero cameras, processing with YOLOv8 and OpenCV overlays, and outputs as annotated MJPEG streams, snapshots, and notifications."
    ))
    maybe_add_image(doc, "flow_input_process_output.png", "Input–Process–Output flow for FalconEye")

    # Add PTZ description
    add_multiline(doc, (
        "PTZ Controller: A simple HTTP-controlled pan mechanism exposed via /camera/pan/<action> enabling left, right, and auto sweeping."
    ))

    # API routes table (as text list)
    routes = extract_routes(BACKEND_FILE) + extract_routes(SIMPLE_BACKEND_FILE)
    if routes:
        add_heading(doc, "API Routes", level=2, center=False)
        for r in routes:
            p = doc.add_paragraph(f"{', '.join(r['methods'])}: {r['path']}")
            apply_tnr(p)
    add_page_break(doc)

    # Chapter 4: Implementation
    add_heading(doc, "Chapter 4: Implementation", level=1)
    impl_points = (
        "- Framework: Flask with CORS, streaming endpoints for MJPEG and snapshot.\n"
        "- Model: ultralytics.YOLO('yolov8s.pt'), inference accelerated via torch.mps on Apple Silicon.\n"
        "- Overlays: Class labels, confidences, FPS, face summaries with adaptive font sizes for mobile.\n"
        "- ESP32 capture loop with shorter timeouts; Pi Zero MJPEG stream decoding with boundary parsing.\n"
        "- Local notifications (with optional FCM legacy/v1 endpoints available)."
    )
    add_multiline(doc, impl_points)

    # Insert code excerpt from backend.py key parts
    backend_excerpt = read_text_if_exists(BACKEND_FILE)
    if backend_excerpt:
        add_heading(doc, "Key Backend Excerpts", level=2, center=False)
        add_code_block(doc, "\n".join(backend_excerpt.splitlines()[170:190] + backend_excerpt.splitlines()[5480:5520]))
    add_page_break(doc)

    # Chapter 5: Testing and Results
    add_heading(doc, "Chapter 5: Testing and Results", level=1)
    add_multiline(doc, (
        "We validated end-to-end flows: camera connectivity, model inference stability, overlay correctness, and Cloudflare remote access. "
        "Test scripts include FCM connectivity checks and camera snapshot/live endpoints validation."
    ))
    # Maybe add class distribution image if present
    maybe_add_image(doc, "class_distribution_present.png", "Representative class distribution of observed objects")
    add_page_break(doc)

    # Chapter 6: Testing
    add_heading(doc, "Chapter 6: Testing", level=1)
    add_multiline(doc, (
        "We evaluated feasibility (economic, technical, social), validated end-to-end flows, and measured accuracy and latency.\n"
        "Figures 6.1–6.3 summarize distribution, precision/recall, and latency trends."
    ))
    maybe_add_image(doc, "testing_class_distribution.png", "Weekly Detections by Class (Figure 6.1)")
    maybe_add_image(doc, "testing_precision_recall.png", "Average Precision/Recall by Class (Figure 6.2)")
    maybe_add_image(doc, "testing_latency_trend.png", "Median Alert Latency Over Time (Figure 6.3)")
    add_page_break(doc)

    # Chapter 7: Discussion
    add_heading(doc, "Chapter 7: Discussion", level=1)
    add_multiline(doc, (
        "FalconEye emphasizes practical edge constraints: unstable Wi‑Fi links, snapshot vs MJPEG trade-offs, and annotation readability. "
        "Cloudflare Tunnel simplifies secure access but adds dependence on third-party uptime."
    ))
    add_page_break(doc)

    # Chapter 8: Conclusion
    add_heading(doc, "Chapter 8: Conclusion", level=1)
    add_multiline(doc, "We presented a deployable, low-latency edge surveillance system with explainable overlays and robust routing.")
    add_page_break(doc)

    # Chapter 9: Future Work
    add_heading(doc, "Chapter 9: Future Work", level=1)
    add_multiline(doc, (
        "- Integrate person re-identification and tracking across cameras.\n"
        "- Add on-device face embedding and privacy-preserving storage.\n"
        "- Offline-first buffering with resumable uploads."
    ))
    add_page_break(doc)

    # References (placeholder)
    add_heading(doc, "References", level=1)
    add_multiline(doc, (
        "[1] Ultralytics YOLOv8 Documentation.\n"
        "[2] OpenCV Documentation.\n"
        "[3] Cloudflare Tunnel Docs."
    ))
    add_page_break(doc)

    # Appendices
    add_heading(doc, "Appendix A: Configuration", level=1)
    for cfg in CONFIG_FILES:
        content = read_text_if_exists(cfg)
        if content:
            p = doc.add_paragraph(os.path.basename(cfg))
            apply_tnr(p, bold=True)
            add_code_block(doc, content)
    add_page_break(doc)

    add_heading(doc, "Appendix B: API Route Inventory", level=1)
    routes_json = json.dumps(routes, indent=2)
    add_code_block(doc, routes_json)
    add_page_break(doc)

    add_heading(doc, "Appendix C: Test Scripts", level=1)
    tests = [
        os.path.join(ROOT, f) for f in os.listdir(ROOT) if f.startswith('test_') and f.endswith('.py')
    ]
    for t in sorted(tests):
        code = read_text_if_exists(t)
        if code:
            p = doc.add_paragraph(os.path.basename(t))
            apply_tnr(p, bold=True)
            add_code_block(doc, code)

    out_path = os.path.expanduser("~/Downloads/VTU_FalconEye_Report_Complete.docx")
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    path = build_document()
    print(f"CREATED {path}")
